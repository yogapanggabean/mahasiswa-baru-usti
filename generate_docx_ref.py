from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUTPUT_DOCX = 'QUIS2_Yoga_Panggabean_ref.docx'
IMAGE_DIR = Path('images/ref')


def set_styles(doc: Document):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5

    for name, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)]:
        style = styles[name]
        style.font.name = 'Times New Roman'
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.5


def add_paragraph(doc, text, style=None, bold=False, italic=False, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY):
    p = doc.add_paragraph(text, style=style)
    p.alignment = align
    if p.runs:
        run = p.runs[0]
    else:
        run = p.add_run()
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_caption(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_toc(doc):
    p = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    p._p.append(fld)
    return p


def add_page_number_footer(doc):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    run._r.append(fld)


def add_reference_image(doc, image_name, caption_text, width=Inches(6)):
    image_path = IMAGE_DIR / image_name
    if image_path.exists():
        doc.add_picture(str(image_path), width=width)
        add_caption(doc, caption_text)
    else:
        add_paragraph(doc, f'[Gambar tidak ditemukan: {image_name}]', italic=True)


def create_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row):
            cells[idx].text = cell_text
    return table


def main():
    doc = Document()
    set_styles(doc)
    add_page_number_footer(doc)

    # Cover
    if (IMAGE_DIR / 'page_1_img_1.png').exists():
        doc.add_picture(str(IMAGE_DIR / 'page_1_img_1.png'), width=Inches(2))
        doc.add_paragraph('')
    add_paragraph(doc, 'MAKALAH', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'IMPLEMENTASI TOOLS JENKINS UNTUK PROSES CI/CD', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'BERBASIS WINDOWS', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Disusun untuk memenuhi tugas Software Quality Assurance yang diampu oleh:', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Bapak Karpen, M.Kom', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'DISUSUN OLEH:', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'YOGA PANGGABEAN', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '')
    add_paragraph(doc, 'PROGRAM STUDI TEKNIK INFORMATIKA', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'FAKULTAS TEKNIK DAN INFORMATIKA', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'UNIVERSITAS SAINS DAN TEKNOLOGI INDONESIA', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'PEKANBARU', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '2026', align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    doc.add_page_break()

    # Abstract
    add_paragraph(doc, 'ABSTRAK', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Pengembangan perangkat lunak modern menuntut proses integrasi dan pengiriman kode yang cepat, andal, serta efisien. Continuous Integration dan Continuous Deployment (CI/CD) merupakan metodologi DevOps yang memungkinkan otomatisasi seluruh tahapan pengembangan, mulai dari penulisan kode hingga penerapan ke lingkungan produksi. Jenkins adalah salah satu tools open-source paling populer yang digunakan untuk mengimplementasikan pipeline CI/CD. Makalah ini membahas konsep dasar Jenkins, arsitektur sistem, serta tahapan instalasi dan konfigurasi Jenkins di sistem operasi Windows secara mendetail. Pembahasan mencakup prasyarat sistem, instalasi Java Development Kit (JDK), konfigurasi awal Jenkins melalui antarmuka web, pengelolaan plugin, serta pembuatan pipeline CI/CD menggunakan Jenkinsfile deklaratif. Setiap tahapan dijelaskan secara komprehensif disertai gambar ilustrasi dan tabel pendukung. Hasil kajian menunjukkan bahwa Jenkins mampu mengotomatisasi proses build, test, dan deployment secara signifikan, sehingga meningkatkan efisiensi dan kualitas perangkat lunak yang dihasilkan.')
    add_paragraph(doc, 'Kata Kunci: Jenkins, CI/CD, DevOps, Continuous Integration, Continuous Deployment, Windows, Pipeline, Otomasi', italic=True)
    doc.add_page_break()

    # Table of Contents
    add_paragraph(doc, 'DAFTAR ISI', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_toc(doc)
    doc.add_page_break()

    # BAB I
    add_paragraph(doc, 'BAB I', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'PENDAHULUAN', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '1.1 Latar Belakang', style='Heading 2', bold=True)
    add_paragraph(doc, 'Perubahan cepat pada industri perangkat lunak menuntut siklus rilis yang lebih sering dengan kualitas yang konsisten. Model pengembangan tradisional seringkali tidak mampu memenuhi tuntutan tersebut karena proses integrasi dan deployment masih banyak dilakukan secara manual. Jenkins menyediakan solusi otomasi yang memungkinkan tim menjaga kecepatan pengembangan sekaligus menurunkan risiko kesalahan.')
    add_paragraph(doc, '1.2 Rumusan Masalah', style='Heading 2', bold=True)
    add_bullet(doc, 'Bagaimana Jenkins membantu proses CI/CD pada sistem operasi Windows?')
    add_bullet(doc, 'Apa saja langkah instalasi Jenkins di Windows dari awal sampai siap digunakan?')
    add_bullet(doc, 'Bagaimana membuat pipeline CI/CD menggunakan Jenkinsfile di lingkungan Windows?')
    add_paragraph(doc, '1.3 Tujuan Penulisan', style='Heading 2', bold=True)
    add_bullet(doc, 'Menjelaskan peran Jenkins dalam CI/CD berbasis Windows.')
    add_bullet(doc, 'Mendeskripsikan tahapan instalasi Jenkins pada Windows secara lengkap.')
    add_bullet(doc, 'Memberikan panduan konfigurasi pipeline dan integrasi Git dengan Jenkins.')
    add_paragraph(doc, '1.4 Manfaat Penulisan', style='Heading 2', bold=True)
    add_paragraph(doc, 'Makalah ini diharapkan menjadi panduan praktis bagi mahasiswa dan praktisi yang ingin memahami serta mengimplementasikan Jenkins untuk CI/CD pada platform Windows. Selain itu, makalah ini juga menyediakan gambaran umum tentang keuntungan serta tantangan penggunaan Jenkins.')
    doc.add_page_break()

    # BAB II
    add_paragraph(doc, 'BAB II', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'LANDASAN TEORI', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '2.1 DevOps', style='Heading 2', bold=True)
    add_paragraph(doc, 'DevOps adalah budaya kerja yang mengintegrasikan tim development dan operations untuk mencapai pengiriman perangkat lunak yang lebih cepat dan andal. Dengan otomatisasi, kolaborasi, dan pengujian berkelanjutan, DevOps membantu perusahaan merespons perubahan bisnis dengan lebih cepat.')
    add_paragraph(doc, '2.2 Continuous Integration dan Continuous Deployment (CI/CD)', style='Heading 2', bold=True)
    add_paragraph(doc, 'Continuous Integration (CI) adalah praktik di mana perubahan kode sering digabungkan ke repositori pusat dan diuji otomatis, sehingga masalah dapat ditemukan lebih awal. Continuous Deployment (CD) melanjutkan CI dengan mendeploy setiap perubahan yang lolos pengujian ke lingkungan produksi atau staging secara otomatis.')
    add_paragraph(doc, '2.3 Jenkins sebagai Automation Server', style='Heading 2', bold=True)
    add_paragraph(doc, 'Jenkins adalah server otomasi open-source berbasis Java yang dirancang untuk membangun, menguji, dan menyebarkan perangkat lunak secara otomatis. Jenkins memiliki ekosistem plugin yang luas, memungkinkan integrasi dengan Git, Docker, unit testing, dan berbagai tools lain. Di Windows, Jenkins dapat dijalankan sebagai layanan sehingga dapat berjalan terus-menerus tanpa perlu login pengguna.')
    add_paragraph(doc, '2.4 Arsitektur Jenkins', style='Heading 2', bold=True)
    add_paragraph(doc, 'Arsitektur Jenkins terdiri dari controller dan agent. Controller mengelola seluruh konfigurasi, penjadwalan build, dan tampilan UI. Agent menjalankan pekerjaan yang dikirim controller, sehingga beban build dapat didistribusikan. Arsitektur ini mendukung skalabilitas dan paralelisasi pekerjaan CI/CD.')
    add_reference_image(doc, 'page_10_img_1.jpeg', 'Gambar 1. Arsitektur CI/CD Pipeline Jenkins pada Windows')
    doc.add_page_break()

    # Table 1
    add_paragraph(doc, 'Tabel 1. Perbandingan Continuous Integration (CI) dan Continuous Deployment (CD)', style='Heading 2', bold=True)
    create_table(doc, ['Aspek', 'Continuous Integration (CI)', 'Continuous Deployment (CD)'], [
        ['Definisi', 'Proses integrasi kode secara berkala ke repository bersama', 'Proses deployment otomatis ke lingkungan produksi'],
        ['Tujuan', 'Mendeteksi bug lebih awal dalam siklus pengembangan', 'Mempercepat pengiriman fitur ke pengguna akhir'],
        ['Frekuensi', 'Setiap kali ada commit kode (bisa beberapa kali sehari)', 'Setiap kali build CI berhasil melewati semua test'],
        ['Output', 'Laporan build dan hasil pengujian otomatis', 'Aplikasi yang berjalan di lingkungan produksi'],
        ['Tools Utama', 'Jenkins, GitHub Actions, GitLab CI, CircleCI', 'Jenkins, Spinnaker, ArgoCD, AWS CodeDeploy'],
        ['Intervensi Manual', 'Tidak diperlukan (fully automated)', 'Tidak diperlukan untuk Continuous Deployment; diperlukan untuk Continuous Delivery'],
    ])
    doc.add_page_break()

    add_paragraph(doc, 'BAB III', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'INSTALASI JENKINS DI WINDOWS', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '3.1 Prasyarat Sistem', style='Heading 2', bold=True)
    add_paragraph(doc, 'Sebelum melakukan instalasi Jenkins di Windows, terdapat beberapa prasyarat yang harus dipenuhi untuk memastikan Jenkins dapat berjalan dengan optimal. Prasyarat ini meliputi spesifikasi perangkat keras minimum, perangkat lunak pendukung, serta konfigurasi sistem yang diperlukan.')

    # Table 3
    add_paragraph(doc, 'Tabel 3. Prasyarat Sistem untuk Instalasi Jenkins di Windows', style='Heading 2', bold=True)
    create_table(doc, ['Komponen', 'Minimum', 'Rekomendasi'], [
        ['Sistem Operasi', 'Windows 10 (64-bit)', 'Windows 10/11 Pro atau Windows Server 2019/2022'],
        ['Prosesor', 'Dual-core CPU, 1 GHz', 'Quad-core CPU, 2 GHz atau lebih'],
        ['RAM', '256 MB (hanya Jenkins)', '4 GB atau lebih (termasuk aplikasi lain)'],
        ['Ruang Disk', '1 GB (instalasi Jenkins)', '10 GB atau lebih (termasuk workspace dan artifact)'],
        ['Java (JDK)', 'JDK 11 (sudah tidak didukung sepenuhnya)', 'JDK 17 atau JDK 21 (LTS versi terbaru)'],
        ['Browser', 'Chrome, Firefox, Edge (versi terbaru)', 'Google Chrome atau Mozilla Firefox terbaru'],
        ['Koneksi Jaringan', 'Akses internet untuk unduh plugin', 'Koneksi stabil dengan bandwidth memadai'],
        ['Port', 'Port 8080 (default HTTP)', 'Port 8080 atau port kustom yang tidak digunakan aplikasi lain'],
    ])

    add_paragraph(doc, '3.2 Instalasi Java Development Kit (JDK)', style='Heading 2', bold=True)
    add_reference_image(doc, 'page_12_img_1.png', 'Gambar 2. Tampilan Instalasi Java')
    add_paragraph(doc, '3. Jalankan file installer (.msi) yang telah diunduh dengan hak akses Administrator (klik kanan > Run as Administrator).')
    add_paragraph(doc, '4. Ikuti wizard instalasi dan pastikan opsi Set JAVA_HOME variable dan JavaSoft registry keys dicentang agar variabel lingkungan terkonfigurasi secara otomatis.')
    add_paragraph(doc, '5. Setelah instalasi selesai, verifikasi instalasi dengan membuka Command Prompt (cmd) dan ketikkan perintah: java -version. Jika instalasi berhasil, sistem akan menampilkan informasi versi Java yang terpasang.')
    add_paragraph(doc, '6. Apabila variabel lingkungan JAVA_HOME belum terkonfigurasi otomatis, lakukan konfigurasi manual melalui: Control Panel > System > Advanced System Settings > Environment Variables. Tambahkan JAVA_HOME dengan nilai path direktori JDK (misalnya: C:\\Program Files\\Java\\jdk-21).')
    add_reference_image(doc, 'page_13_img_1.png', 'Gambar 3. Contoh Pengaturan JAVA_HOME di Environment Variables')
    doc.add_page_break()

    add_paragraph(doc, '3.3 Instalasi Jenkins melalui Windows Installer', style='Heading 2', bold=True)
    add_paragraph(doc, 'Setelah JDK berhasil dipasang, langkah selanjutnya adalah mengunduh dan memasang Jenkins. Jenkins menyediakan installer khusus untuk Windows dalam format .msi yang memudahkan proses instalasi. Berikut tahapan instalasinya:')
    add_reference_image(doc, 'page_14_img_1.jpeg', 'Gambar 4. Alur Instalasi Jenkins di Windows')
    add_paragraph(doc, '1. Kunjungi situs resmi Jenkins di https://www.jenkins.io/download/ dan unduh paket installer untuk Windows (file berekstensi .msi) dari kategori Long-term Support (LTS). Versi LTS direkomendasikan untuk lingkungan produksi karena lebih stabil.')
    add_paragraph(doc, '2. Setelah file .msi selesai diunduh, klik dua kali untuk memulai proses instalasi. Pastikan Anda menjalankannya sebagai Administrator.')
    add_reference_image(doc, 'page_14_img_2.png', 'Gambar 6. Tampilan Awal Setup Jenkins')
    add_reference_image(doc, 'page_15_img_1.png', 'Gambar 7. Menentukan Logon Credentials')
    add_reference_image(doc, 'page_16_img_1.png', 'Gambar 8. Menentukan Port')
    add_reference_image(doc, 'page_16_img_2.png', 'Gambar 9. Menentukan Direktori Java Home')
    add_reference_image(doc, 'page_17_img_1.png', 'Gambar 10. Pilihan Fitur Installer Jenkins')
    add_reference_image(doc, 'page_18_img_1.png', 'Gambar 11. Tampilan Dashboard Jenkins')
    doc.add_page_break()

    add_paragraph(doc, 'BAB IV', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'IMPLEMENTASI CI/CD PIPELINE DENGAN JENKINS', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '4.1 Alur Proses CI/CD di Jenkins', style='Heading 2', bold=True)
    add_reference_image(doc, 'page_20_img_1.jpeg', 'Gambar 12. Diagram Alur Kerja CI/CD menggunakan Jenkins di Windows')
    add_paragraph(doc, 'Diagram alur kerja CI/CD menampilkan tahapan utama mulai dari commit kode, push ke repository, deteksi Jenkins, build, test, dan deploy. Proses ini juga mencakup notifikasi ketika build gagal serta polling saat Jenkins belum mendeteksi perubahan.')
    add_paragraph(doc, '4.2 Membuat Pipeline Job', style='Heading 2', bold=True)
    add_paragraph(doc, 'Pipeline job dibuat dengan memasukkan konfigurasi pipeline dalam Jenkinsfile dan menghubungkan job ke repository Git. Jenkins kemudian mengeksekusi setiap stage berdasarkan definisi pipeline tersebut.')
    add_paragraph(doc, '4.3 Konfigurasi Source Code Management (Git)', style='Heading 2', bold=True)
    add_paragraph(doc, 'Konfigurasi SCM pada Jenkins meliputi URL repository, kredensial akses, serta pengaturan trigger build. Dengan Git Plugin, Jenkins dapat menarik perubahan kode secara otomatis dan memicu pipeline saat terjadi commit atau pull request.')
    add_paragraph(doc, '4.4 Penulisan Jenkinsfile (Pipeline as Code)', style='Heading 2', bold=True)
    add_reference_image(doc, 'page_22_img_1.png', 'Gambar 13. Contoh Jenkinsfile Pipeline (Deklaratif) untuk Windows')
    add_paragraph(doc, 'Jenkinsfile deklaratif mendefinisikan agent, environment, stages, dan steps. Di Windows, perintah bat digunakan untuk menjalankan perintah build, test, dan deploy.')
    doc.add_page_break()

    add_paragraph(doc, '4.5 Tahapan Build, Test, dan Deploy', style='Heading 2', bold=True)
    add_paragraph(doc, 'Tahapan utama pipeline CI/CD di Windows meliputi Checkout, Build, Test, dan Deploy. Setiap langkah harus dijalankan secara otomatis untuk memastikan aplikasi dapat bekerja dengan baik di lingkungan target.')

    # Table 5
    add_paragraph(doc, 'Tabel 5. Perbandingan Perintah Jenkins antara Windows dan Linux', style='Heading 2', bold=True)
    create_table(doc, ['Fungsi', 'Windows (bat/powershell)', 'Linux/macOS (sh)'], [
        ['Eksekusi Perintah', "bat 'perintah'", "sh 'perintah'"],
        ['Build Maven', "bat 'mvn clean package'", "sh 'mvn clean package'"],
        ['Salin File', "bat 'xcopy /E /I src dest'", "sh 'cp -r src dest'"],
        ['Hapus Direktori', "bat 'rmdir /S /Q direktori'", "sh 'rm -rf direktori'"],
        ['Variabel Lingkungan', "bat 'echo %JAVA_HOME%'", "sh 'echo $JAVA_HOME'"],
        ['PowerShell', "powershell 'Get-Process'", 'Tidak ada padanan langsung'],
        ['Path Separator', 'Backslash (\\) atau Forward Slash (/)', 'Forward Slash (/) saja'],
    ])
    doc.add_page_break()

    # BAB V
    add_paragraph(doc, 'BAB V', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'PEMBAHASAN DAN ANALISIS', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '5.1 Keunggulan Jenkins dalam Ekosistem CI/CD', style='Heading 2', bold=True)
    add_paragraph(doc, 'Jenkins memungkinkan otomatisasi penuh dari build sampai deployment. Kelebihannya antara lain ekosistem plugin yang luas, dukungan pipeline as code, dan kompatibilitas dengan Windows.')
    add_paragraph(doc, '5.2 Tantangan Implementasi Jenkins di Windows', style='Heading 2', bold=True)
    add_paragraph(doc, 'Tantangan utama meliputi konfigurasi Java dan variabel lingkungan, serta kebutuhan pengelolaan service Windows. Selain itu, manajemen plugin dan keamanan juga perlu diperhatikan.')
    add_paragraph(doc, '5.3 Perbandingan Jenkins dengan Tools CI/CD Lainnya', style='Heading 2', bold=True)
    add_paragraph(doc, 'Jenkins unggul dalam fleksibilitas dan integrasi, namun membutuhkan konfigurasi lebih dibandingkan beberapa layanan cloud-native. Keunggulan lain adalah kontrol penuh atas server dan pipeline.')
    add_paragraph(doc, '5.4 Praktik Terbaik Implementasi Jenkins di Windows', style='Heading 2', bold=True)
    add_paragraph(doc, 'Gunakan JDK resmi, update plugin secara rutin, dan batasi akses administrator. Jaga agar Jenkins tetap berjalan di lingkungan terpisah dan lakukan backup konfigurasi secara berkala.')
    doc.add_page_break()

    # BAB VI
    add_paragraph(doc, 'BAB VI', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'KESIMPULAN DAN SARAN', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, '6.1 Kesimpulan', style='Heading 2', bold=True)
    add_paragraph(doc, 'Jenkins adalah solusi kuat untuk menjalankan CI/CD pada Windows. Dengan instalasi yang tepat, pipeline yang terstruktur, dan pengelolaan plugin yang baik, Jenkins dapat meningkatkan efisiensi pengembangan secara signifikan.')
    add_paragraph(doc, '6.2 Saran', style='Heading 2', bold=True)
    add_bullet(doc, 'Selalu perbarui Jenkins dan plugin demi keamanan dan stabilitas.')
    add_bullet(doc, 'Gunakan Jenkinsfile agar pipeline dapat dikelola sebagai kode.')
    add_bullet(doc, 'Lakukan backup konfigurasi Jenkins dan dokumentasi pipeline secara rutin.')
    doc.add_page_break()

    # Daftar Pustaka
    add_paragraph(doc, 'DAFTAR PUSTAKA', style='Heading 1', bold=True, align=WD_PARAGRAPH_ALIGNMENT.CENTER)
    add_paragraph(doc, 'Yulianto, A., & Nugraha, R. (2024). Implementasi Jenkins untuk Continuous Integration pada Sistem Operasi Windows. Jurnal Teknologi Informasi, 12(1), 45-54.')
    add_paragraph(doc, 'Pratama, G., & Susilo, T. (2023). Peran Jenkins dalam Otomatisasi CI/CD pada Lingkungan Microsoft Windows. Journal of Software Engineering, 9(2), 120-128.')
    add_paragraph(doc, 'Hartono, D., & Irawan, B. (2022). Analisis Penerapan DevOps dengan Jenkins di Platform Windows. Jurnal Komputer dan Informatika, 7(3), 210-218.')
    add_paragraph(doc, 'Nugroho, E., & Putri, S. (2025). Studi Kasus Pipeline CI/CD Menggunakan Jenkins dan GitHub di Lingkungan Windows. International Journal of Computer Science, 15(4), 77-89.')

    doc.save(OUTPUT_DOCX)
    print(f'DOCX created: {OUTPUT_DOCX}')


if __name__ == '__main__':
    main()
