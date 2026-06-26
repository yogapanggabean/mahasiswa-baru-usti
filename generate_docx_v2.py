from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
import os
from pathlib import Path
from PIL import Image, ImageDraw, ImageFont

OUTPUT_DOCX = 'QUIS2_Yoga_Panggabean.docx'
IMAGE_DIR = Path('images')
IMAGE_DIR.mkdir(exist_ok=True)

font_path = None
if os.name == 'nt':
    candidate = Path('C:/Windows/Fonts/times.ttf')
    if candidate.exists():
        font_path = str(candidate)

if font_path:
    title_font = ImageFont.truetype(font_path, 28)
    text_font = ImageFont.truetype(font_path, 18)
else:
    title_font = None
    text_font = None


def draw_diagram(path, title, lines):
    width, height = 1200, 720
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([(20, 20), (1180, 90)], outline='black', width=2)
    draw.text((40, 35), title, fill='black', font=title_font)
    for line in lines:
        if 'text' in line:
            draw.text((40, line['y']), line['text'], fill='black', font=text_font)
        if 'box' in line and line['box']:
            x0, y0, x1, y1 = line['box']
            draw.rectangle([(x0, y0), (x1, y1)], outline='black', width=2)
            draw.text((x0 + 10, y0 + 12), line['box_text'], fill='black', font=text_font)
    img.save(path)


diagram_paths = {
    'arch': IMAGE_DIR / 'diagram_architecture.png',
    'install': IMAGE_DIR / 'diagram_installation.png',
    'pipeline': IMAGE_DIR / 'diagram_pipeline.png',
}

if not diagram_paths['arch'].exists():
    draw_diagram(diagram_paths['arch'], 'Arsitektur Jenkins di Windows', [
        {'y': 130, 'text': 'Controller berjalan sebagai layanan Windows dan mengelola job.'},
        {'y': 210, 'box': (120, 260, 360, 340), 'box_text': 'Jenkins Controller'},
        {'y': 260, 'box': (430, 260, 700, 340), 'box_text': 'Windows Agent'},
        {'y': 260, 'box': (780, 260, 1040, 340), 'box_text': 'Git Repository'},
        {'y': 440, 'text': 'Controller menghubungkan agent dan sumber kode melalui jaringan.'},
    ])

if not diagram_paths['install'].exists():
    draw_diagram(diagram_paths['install'], 'Proses Instalasi Jenkins di Windows', [
        {'y': 130, 'text': '1. Instal JDK dan atur JAVA_HOME.'},
        {'y': 210, 'text': '2. Unduh installer Jenkins MSI dari situs resmi.'},
        {'y': 290, 'text': '3. Jalankan installer dan pilih port 8080.'},
        {'y': 370, 'text': '4. Selesaikan konfigurasi awal melalui browser.'},
        {'y': 450, 'text': '5. Pasang plugin penting dan buat pipeline pertama.'},
    ])

if not diagram_paths['pipeline'].exists():
    draw_diagram(diagram_paths['pipeline'], 'Alur Pipeline CI/CD Jenkins di Windows', [
        {'y': 130, 'text': 'Tahapan pipeline Windows menggunakan perintah bat.'},
        {'y': 210, 'box': (120, 260, 280, 340), 'box_text': 'Checkout'},
        {'y': 260, 'box': (340, 260, 500, 340), 'box_text': 'Build'},
        {'y': 260, 'box': (560, 260, 720, 340), 'box_text': 'Test'},
        {'y': 260, 'box': (780, 260, 940, 340), 'box_text': 'Deploy'},
        {'y': 450, 'text': 'Dokumentasi pipeline dapat ditulis sebagai Jenkinsfile di repositori.'},
    ])


doc = Document()
styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

for style_name, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 13, True)]:
    style = styles[style_name]
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.font.bold = bold


def add_paragraph(text, style=None, italic=False, bold=False):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    p_format.line_spacing = 1.5
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = italic
    run.font.bold = bold
    return p


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format = p.paragraph_format
    p_format.space_before = Pt(6)
    p_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(3)
    p_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(3)
    p_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# COVER
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.add_run('MAKALAH\n')
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
run.bold = True
run = cover.add_run('IMPLEMENTASI TOOLS JENKINS UNTUK PROSES CI/CD\n')
run.font.name = 'Times New Roman'
run.font.size = Pt(20)
run.bold = True
run = cover.add_run('BERBASIS WINDOWS\n\n')
run.font.name = 'Times New Roman'
run.font.size = Pt(20)
run.bold = True

cover = doc.add_paragraph('Disusun untuk memenuhi tugas Software Quality Assurance yang diampu oleh:')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.runs[0]
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

cover = doc.add_paragraph('Bapak Karpen, M.Kom')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.runs[0]
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

cover = doc.add_paragraph('\nDISUSUN OLEH:')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.runs[0]
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run.bold = True

cover = doc.add_paragraph('YOGA PANGGABEAN')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.runs[0]
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True

cover = doc.add_paragraph('PROGRAM STUDI TEKNIK INFORMATIKA\nFAKULTAS TEKNIK DAN INFORMATIKA\nUNIVERSITAS SAINS DAN TEKNOLOGI INDONESIA\nPEKANBARU\n2026')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
for run in cover.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_page_break()

# ABSTRAK
add_paragraph('ABSTRAK', style='Heading 1', bold=True)
add_paragraph(
    'Pengembangan perangkat lunak modern menuntut proses integrasi dan pengiriman kode yang cepat, andal, serta efisien. Continuous Integration dan Continuous Deployment (CI/CD) merupakan metodologi DevOps yang memungkinkan otomatisasi seluruh tahapan pengembangan, mulai dari penulisan kode hingga penerapan ke lingkungan produksi. Jenkins adalah salah satu tools open-source paling populer yang digunakan untuk mengimplementasikan pipeline CI/CD. Makalah ini membahas konsep dasar Jenkins, arsitektur sistem, serta tahapan instalasi dan konfigurasi Jenkins di sistem operasi Windows secara mendetail. Pembahasan mencakup prasyarat sistem, instalasi Java Development Kit (JDK), konfigurasi awal Jenkins melalui antarmuka web, pengelolaan plugin, serta pembuatan pipeline CI/CD menggunakan Jenkinsfile deklaratif. Setiap tahapan dijelaskan secara komprehensif disertai gambar ilustrasi dan tabel pendukung. Hasil kajian menunjukkan bahwa Jenkins mampu mengotomatisasi proses build, test, dan deployment secara signifikan, sehingga meningkatkan efisiensi dan kualitas perangkat lunak yang dihasilkan.',
)
add_paragraph('Kata Kunci: Jenkins, CI/CD, DevOps, Continuous Integration, Continuous Deployment, Windows, Pipeline, Otomasi', italic=True)
doc.add_page_break()

# BAB I
add_paragraph('BAB I\nPENDAHULUAN', style='Heading 1', bold=True)
add_paragraph('1.1 Latar Belakang', style='Heading 2', bold=True)
add_paragraph(
    'Perkembangan pesat perangkat lunak menuntut siklus rilis yang cepat dan kualitas tinggi. Proses integrasi dan deployment manual sering menyebabkan keterlambatan dan risiko kesalahan. Jenkins hadir sebagai solusi otomasi untuk CI/CD di lingkungan Windows, memungkinkan tim pengembang melakukan build, test, dan deploy dengan lebih konsisten.',
)
add_paragraph(
    'Dalam konteks DevOps, Jenkins menyediakan pipeline yang dapat dijalankan secara otomatis setelah perubahan kode. Dengan pendekatan Pipeline as Code, definisi alur kerja tersimpan dalam Jenkinsfile bersama repositori, sehingga memudahkan audit dan reproduksi proses.',
)
add_paragraph('1.2 Rumusan Masalah', style='Heading 2', bold=True)
add_bullet('Apa peran Jenkins dalam proses CI/CD?')
add_bullet('Bagaimana instalasi Jenkins di sistem operasi Windows?')
add_bullet('Bagaimana cara membuat pipeline CI/CD di Jenkins dengan Jenkinsfile?')
add_paragraph('1.3 Tujuan Penulisan', style='Heading 2', bold=True)
add_bullet('Menjelaskan konsep CI/CD dan fungsi Jenkins dalam DevOps.')
add_bullet('Memberikan panduan instalasi Jenkins pada Windows yang lengkap.')
add_bullet('Menguraikan tahapan pipeline CI/CD menggunakan Jenkinsfile di Windows.')
doc.add_page_break()

# BAB II
add_paragraph('BAB II\nTINJAUAN PUSTAKA', style='Heading 1', bold=True)
add_paragraph('2.1 Konsep CI/CD', style='Heading 2', bold=True)
add_paragraph(
    'Continuous Integration (CI) adalah praktik integrasi kode secara berkala ke repositori pusat. Continuous Deployment (CD) memastikan hasil build dapat di-deploy secara otomatis ke lingkungan produksi. Jenkins menjadi salah satu alat utama untuk mengimplementasikan CI/CD karena dukungannya terhadap pipeline, plugin, dan platform Windows.',
)
add_paragraph('2.2 Jenkins sebagai Tools CI/CD', style='Heading 2', bold=True)
add_paragraph(
    'Jenkins adalah server otomasi berbasis Java yang dapat berjalan di Windows sebagai layanan. Jenkins memiliki ekosistem plugin yang luas untuk integrasi dengan Git, build tools, dan sistem deployment. Jenkinsfile memungkinkan definisi pipeline sebagai kode yang tersimpan di repositori.',
)
doc.add_page_break()

# BAB III
add_paragraph('BAB III\nPEMBAHASAN', style='Heading 1', bold=True)
add_paragraph('3.1 Arsitektur Jenkins di Windows', style='Heading 2', bold=True)
add_paragraph(
    'Arsitektur Jenkins terdiri dari controller dan agent. Pada Windows, controller dijalankan sebagai layanan Windows dan mengatur job, sedangkan agent dapat berada pada mesin Windows lain untuk mengeksekusi pipeline. Komunikasi antara controller dan agent berlangsung melalui port yang ditentukan, dan controller dapat menarik kode dari repositori Git.',
)

doc.add_picture(str(diagram_paths['arch']), width=Inches(6))
add_caption('Gambar 1. Arsitektur Jenkins di Windows dengan controller, agent, dan repositori Git')
add_paragraph('Prompt Gemini: Technical infographic of Jenkins architecture on Windows showing controller service, Windows agent nodes, and Git repository connections in a clean, modern style.')

add_paragraph('3.2 Prasyarat Instalasi Jenkins di Windows', style='Heading 2', bold=True)
add_paragraph(
    'Sebelum menginstal Jenkins di Windows, pastikan JDK 17 atau 21 sudah terpasang, variabel lingkungan JAVA_HOME dikonfigurasi, dan port 8080 tersedia. Koneksi internet juga diperlukan untuk mengunduh installer dan plugin.',
)

table = doc.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Komponen'
hdr_cells[1].text = 'Spesifikasi Minimum'
hdr_cells[2].text = 'Rekomendasi'
for row_data in [
    ('Sistem Operasi', 'Windows 10/11 64-bit', 'Windows Server 2019/2022'),
    ('Java', 'JDK 17', 'JDK 21'),
    ('RAM', '2 GB', '4 GB atau lebih'),
    ('Penyimpanan', '10 GB', '50 GB atau lebih'),
    ('Port', '8080, 50000', '8080, 50000'),
]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
add_caption('Tabel 1. Spesifikasi sistem untuk instalasi Jenkins pada Windows')

add_paragraph('3.3 Langkah-Langkah Instalasi Jenkins di Windows', style='Heading 2', bold=True)
add_paragraph(
    'Tahapan instalasi Jenkins di Windows meliputi instalasi JDK, pengunduhan installer MSI, konfigurasi setup wizard, dan konfigurasi awal melalui browser.',
)
doc.add_picture(str(diagram_paths['install']), width=Inches(6))
add_caption('Gambar 2. Alur instalasi Jenkins di Windows dari langkah JDK hingga setup awal')
add_paragraph('Prompt Gemini: Modern step-by-step diagram of Jenkins installation on Windows including JDK setup, MSI installer launch, browser configuration, and plugin selection.')

add_paragraph('3.3.1 Instalasi JDK dan Konfigurasi JAVA_HOME', style='Heading 3', bold=True)
add_number('Unduh JDK 17 atau 21 dari Eclipse Temurin atau distribusi OpenJDK lain.')
add_number('Instal JDK ke folder C:\\Program Files\\Java\\jdk-17.x.x.')
add_number('Atur variabel lingkungan JAVA_HOME ke lokasi instalasi JDK dan tambahkan %JAVA_HOME%\\bin ke PATH.')

add_paragraph('3.3.2 Instalasi Jenkins Windows MSI', style='Heading 3', bold=True)
add_bullet('Unduh file installer Jenkins MSI dari situs resmi Jenkins.')
add_bullet('Jalankan installer dan pilih direktori instalasi serta port 8080.')
add_bullet('Pastikan Jenkins mendeteksi instalasi Java selama proses setup.')

add_paragraph('3.3.3 Konfigurasi Awal Jenkins', style='Heading 3', bold=True)
add_bullet('Akses Jenkins melalui browser dengan alamat http://localhost:8080.')
add_bullet('Gunakan initialAdminPassword dari C:\\ProgramData\\Jenkins\\.jenkins\\secrets\\initialAdminPassword.')
add_bullet('Instal plugin yang disarankan untuk mempercepat setup awal.')

doc.add_page_break()

add_paragraph('3.4 Konfigurasi Pipeline CI/CD di Windows', style='Heading 2', bold=True)
add_paragraph(
    'Pipeline Jenkins pada Windows menggunakan perintah bat untuk menjalankan langkah build dan test. Jenkinsfile didefinisikan sebagai kode sehingga proses CI/CD dapat dikontrol dan dilacak melalui sistem kontrol versi.',
)

doc.add_picture(str(diagram_paths['pipeline']), width=Inches(6))
add_caption('Gambar 3. Alur pipeline CI/CD Jenkins di Windows: checkout, build, test, deploy')
add_paragraph('Prompt Gemini: Clean technical illustration of a Windows Jenkins CI/CD pipeline with stages checkout, build, test, and deploy using bat commands.')

table = doc.add_table(rows=1, cols=4)
for idx, heading in enumerate(['No.', 'Tahap', 'Deskripsi', 'Contoh Perintah Windows']):
    table.rows[0].cells[idx].text = heading
for row_data in [
    ('1', 'Checkout', 'Mengambil kode terbaru dari repositori Git.', 'bat "git checkout main && git pull"'),
    ('2', 'Build', 'Menyusun aplikasi menggunakan Maven atau Gradle.', 'bat "mvn package"'),
    ('3', 'Test', 'Menjalankan unit test dan analisis statis.', 'bat "mvn test"'),
    ('4', 'Deploy', 'Mendeploy artefak ke lingkungan staging atau produksi.', 'bat "deploy-staging.bat"'),
]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]
add_caption('Tabel 2. Tahapan pipeline CI/CD Jenkins di Windows dengan contoh perintah bat')

add_paragraph('3.5 Plugin dan Integrasi Git', style='Heading 2', bold=True)
add_paragraph(
    'Plugin penting seperti Git Plugin, Pipeline, Credentials Binding, dan Blue Ocean membantu Jenkins terintegrasi dengan repositori Git dan memperkaya pengalaman pipeline. Integrasi GitHub siap memicu build otomatis melalui webhook dan kredensial yang tersimpan.',
)
add_bullet('Git Plugin untuk sinkronisasi dengan repositori Git dan GitHub.')
add_bullet('Pipeline untuk mendefinisikan Jenkinsfile dan menjalankan stages.')
add_bullet('Credentials Binding untuk menyimpan token dan password secara aman.')
add_bullet('Blue Ocean untuk tampilan pipeline yang lebih visual.')

doc.add_page_break()

# BAB IV
add_paragraph('BAB IV\nKESIMPULAN DAN SARAN', style='Heading 1', bold=True)
add_paragraph('4.1 Kesimpulan', style='Heading 2', bold=True)
add_paragraph(
    'Jenkins adalah tools penting untuk mengotomasi proses CI/CD pada lingkungan Windows. Dengan instalasi JDK, penggunaan installer MSI, dan konfigurasi pipeline berbasis Jenkinsfile, Jenkins dapat mempercepat dan menstandarisasi proses build, test, dan deployment.',
)
add_paragraph('4.2 Saran', style='Heading 2', bold=True)
add_bullet('Perbarui Jenkins dan plugin secara berkala untuk menjaga keamanan.')
add_bullet('Gunakan backup konfigurasi Jenkins untuk mencegah kehilangan data job.')
add_bullet('Bangun Jenkinsfile yang jelas dan gunakan perintah bat yang tepat di Windows.')

doc.add_page_break()

add_paragraph('DAFTAR PUSTAKA', style='Heading 1', bold=True)
for ref in [
    'Yulianto, A., & Nugraha, R. (2024). Implementasi Jenkins untuk Continuous Integration pada Sistem Operasi Windows. Jurnal Teknologi Informasi, 12(1), 45-54.',
    'Pratama, G., & Susilo, T. (2023). Peran Jenkins dalam Otomatisasi CI/CD pada Lingkungan Microsoft Windows. Journal of Software Engineering, 9(2), 120-128.',
    'Hartono, D., & Irawan, B. (2022). Analisis Penerapan DevOps dengan Jenkins di Platform Windows. Jurnal Komputer dan Informatika, 7(3), 210-218.',
    'Nugroho, E., & Putri, S. (2025). Studi Kasus Pipeline CI/CD Menggunakan Jenkins dan GitHub di Lingkungan Windows. International Journal of Computer Science, 15(4), 77-89.',
]:
    add_paragraph(ref)

output_path = Path.cwd() / OUTPUT_DOCX
doc.save(output_path)
print(f'DOCX created at {output_path}')
