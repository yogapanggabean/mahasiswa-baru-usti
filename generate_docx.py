from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image, ImageDraw, ImageFont
import os

OUTPUT_DOCX = 'QUIS2.docx'

# Create diagram images programmatically
os.makedirs('images', exist_ok=True)

font_path = None
try:
    from pathlib import Path
    if os.name == 'nt':
        font_path = str(Path('C:/Windows/Fonts/times.ttf'))
    else:
        font_path = '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'
except Exception:
    font_path = None

if font_path and os.path.exists(font_path):
    title_font = ImageFont.truetype(font_path, 24)
    text_font = ImageFont.truetype(font_path, 18)
else:
    title_font = None
    text_font = None


def draw_diagram(path, title, lines):
    width, height = 1200, 700
    bg = (255, 255, 255)
    img = Image.new('RGB', (width, height), color=bg)
    draw = ImageDraw.Draw(img)
    draw.rectangle([(20, 20), (1180, 80)], outline='black', width=2)
    draw.text((40, 30), title, fill='black', font=title_font)
    for line in lines:
        if 'text' in line:
            draw.text((40, line['y']), line['text'], fill='black', font=text_font)
        if 'box' in line and line['box']:
            x0, y0, x1, y1 = line['box']
            draw.rectangle([(x0, y0), (x1, y1)], outline='black', width=2)
            draw.text((x0 + 10, y0 + 10), line['box_text'], fill='black', font=text_font)
    img.save(path)


# Create three diagrams if not exist
if not os.path.exists('images/diagram_architecture.png'):
    draw_diagram('images/diagram_architecture.png', 'Arsitektur Jenkins pada Windows', [
        {'y': 120, 'text': 'Controller Jenkins berjalan sebagai layanan Windows di port 8080.'},
        {'y': 200, 'box': (120, 260, 360, 340), 'box_text': 'Jenkins Controller'},
        {'y': 330, 'box': (420, 260, 700, 340), 'box_text': 'Node/Agent Windows'},
        {'y': 330, 'box': (760, 260, 1040, 340), 'box_text': 'Git Repository / Source Control'},
        {'y': 430, 'text': 'Controller berkomunikasi dengan Agent melalui port 50000 dan menarik kode dari Git.'}
    ])

if not os.path.exists('images/diagram_installation.png'):
    draw_diagram('images/diagram_installation.png', 'Tahapan Instalasi Jenkins di Windows', [
        {'y': 120, 'text': '1. Instalasi JDK 17/21 dan konfigurasi JAVA_HOME.'},
        {'y': 200, 'text': '2. Unduh Jenkins Windows MSI dari situs resmi.'},
        {'y': 280, 'text': '3. Jalankan installer, pilih port 8080 dan layanan Windows.'},
        {'y': 360, 'text': '4. Akses http://localhost:8080 dan selesaikan konfigurasi awal.'},
        {'y': 440, 'text': '5. Instal plugin penting dan buat pipeline CI/CD pertama.'}
    ])

if not os.path.exists('images/diagram_pipeline.png'):
    draw_diagram('images/diagram_pipeline.png', 'Alur Pipeline CI/CD Jenkins di Windows', [
        {'y': 120, 'text': 'Stage utama pada pipeline Windows: Checkout, Build, Test, Paket, Deploy.'},
        {'y': 200, 'box': (120, 260, 320, 340), 'box_text': 'Checkout'},
        {'y': 330, 'box': (380, 260, 560, 340), 'box_text': 'Build'},
        {'y': 330, 'box': (620, 260, 800, 340), 'box_text': 'Test'},
        {'y': 330, 'box': (860, 260, 1040, 340), 'box_text': 'Deploy'},
        {'y': 430, 'text': 'Pipeline menggunakan perintah bat untuk eksekusi di Windows.'}
    ])

# Build document
doc = Document()
styles = doc.styles
normal_style = styles['Normal']
normal_font = normal_style.font
normal_font.name = 'Times New Roman'
normal_font.size = Pt(12)

# Create heading styles if not exist
for style_name, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 13, True)]:
    style = styles[style_name]
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.font.bold = bold


def add_paragraph(text, style=None, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, italic=False, bold=False):
    p = doc.add_paragraph(text, style=style)
    p.alignment = align
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.italic = italic
    run.font.bold = bold
    return p


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

# Cover
cover = doc.add_paragraph()
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = cover.add_run('MAKALAH\n')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.bold = True
run = cover.add_run('IMPLEMENTASI TOOLS JENKINS UNTUK PROSES CI/CD BERBASIS WINDOWS\n')
run.font.name = 'Times New Roman'
run.font.size = Pt(18)
run.bold = True
cover = doc.add_paragraph('')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
cover = doc.add_paragraph('2025', style='Normal')
cover.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Page break
doc.add_page_break()

# Abstract
doc.add_heading('ABSTRAK', level=1)
add_paragraph(
    'Jenkins adalah perangkat lunak open-source yang sangat populer untuk mengotomasi praktik Continuous Integration dan Continuous Delivery (CI/CD). Makalah ini membahas instalasi Jenkins di Windows, konfigurasi pipeline Jenkinsfile dengan perintah bat, dan pengaturan otomatisasi build. Fokus pembahasan adalah integrasi Jenkins dengan repositori Git, instalasi plugin penting, serta best practice keamanan pada lingkungan Windows.',
    style=None
)
add_paragraph('Kata Kunci: Jenkins, CI/CD, Windows, Instalasi, Pipeline, Automation', italic=True)

doc.add_page_break()

doc.add_heading('BAB I\nPENDAHULUAN', level=1)
doc.add_heading('1.1 Latar Belakang', level=2)
add_paragraph(
    'Perkembangan pesat perangkat lunak menuntut siklus pengembangan yang lebih cepat dan kualitas lebih baik. Praktik manual pada integrasi dan deployment seringkali menyebabkan delay dan kesalahan yang dapat merugikan pengembang dan pemangku kepentingan. Jenkins hadir sebagai alat otomasi untuk mengatasi tantangan tersebut dengan mengintegrasikan proses build, test, dan deploy secara otomatis.',
)
add_paragraph(
    'Dalam konteks DevOps, Jenkins mendukung CI/CD dengan memanfaatkan plugin ekstensif dan kemampuan pipeline as code. Di Windows, Jenkins bekerja bersama JDK dan layanan Windows untuk menyediakan server CI yang dapat diakses melalui browser.',
)

doc.add_heading('1.2 Rumusan Masalah', level=2)
add_bullet('Apa peran Jenkins dalam proses CI/CD?')
add_bullet('Bagaimana instalasi Jenkins di Windows secara lengkap?')
add_bullet('Bagaimana mengkonfigurasi pipeline CI/CD dengan Jenkinsfile pada Windows?')

doc.add_heading('1.3 Tujuan Penulisan', level=2)
add_bullet('Menjelaskan konsep CI/CD dan fungsi Jenkins dalam ekosistem DevOps.')
add_bullet('Memberikan panduan instalasi Jenkins pada sistem operasi Windows.')
add_bullet('Menguraikan tahapan pipeline CI/CD menggunakan Jenkinsfile dengan perintah bat.')

doc.add_page_break()

doc.add_heading('BAB II\nTINJAUAN PUSTAKA', level=1)
doc.add_heading('2.1 Konsep CI/CD', level=2)
add_paragraph(
    'Continuous Integration (CI) adalah praktik di mana pengembang menggabungkan perubahan kode secara teratur ke repositori shared, sehingga memungkinkan build dan pengujian otomatis untuk mendeteksi bug lebih awal. Continuous Delivery/Deployment (CD) memperluas CI dengan menyiapkan artefak yang siap dirilis dan dapat dideploy secara otomatis.',
)
add_paragraph(
    'Jenkins mendukung kedua konsep tersebut melalui pipeline yang dapat dijalankan setelah setiap perubahan kode. Menurut studi dalam empat tahun terakhir, penerapan pipeline CI/CD dapat mempercepat rilis perangkat lunak hingga 50 persen dan mengurangi kesalahan integrasi.',
)

doc.add_heading('2.2 Jenkins sebagai Tools CI/CD', level=2)
add_paragraph(
    'Jenkins adalah server otomasi berbasis Java yang digunakan untuk membangun, menguji, dan mengirimkan perangkat lunak secara otomatis. Jenkins menawarkan lebih dari 1.800 plugin yang memungkinkan integrasi dengan sistem kontrol versi, alat build, platform cloud, dan sistem notifikasi.',
)
add_paragraph('Keunggulan Jenkins antara lain fleksibilitas plugin, dukungan platform Windows, dan kemampuan pipeline as code melalui Jenkinsfile.')

doc.add_page_break()

doc.add_heading('BAB III\nPEMBAHASAN', level=1)
doc.add_heading('3.1 Arsitektur Jenkins di Windows', level=2)
add_paragraph(
    'Arsitektur Jenkins umumnya terdiri dari controller (master) dan agent (node). Pada Windows, controller berjalan sebagai layanan Windows dan mengendalikan eksekusi job. Agent dapat dijalankan di mesin Windows atau platform lain untuk mendistribusikan beban build.',
)
add_paragraph(
    'Controller berkomunikasi dengan agen melalui jaringan, mengambil kode sumber dari repositori Git, dan mengelola pipeline serta history build.',
)
doc.add_picture('images/diagram_architecture.png', width=Inches(6))
add_caption('Gambar 1. Arsitektur Jenkins pada lingkungan Windows dengan controller, agent, dan repositori Git')

doc.add_heading('3.2 Prasyarat Instalasi Jenkins di Windows', level=2)
add_paragraph(
    'Sebelum instalasi, pastikan sistem memenuhi persyaratan minimal: Windows 10/Server 2019 atau lebih baru, JDK 17/21 terinstal, minimal 2 GB RAM, dan port 8080 tidak digunakan. Koneksi internet diperlukan untuk mengunduh installer dan plugin.',
)

table = doc.add_table(rows=1, cols=3)
for idx, heading in enumerate(['Komponen', 'Spesifikasi Minimum', 'Rekomendasi']):
    cell = table.rows[0].cells[idx]
    cell.text = heading
for row_data in [
    ('Sistem Operasi', 'Windows 10/11 64-bit', 'Windows Server 2019/2022'),
    ('Java', 'JDK 17', 'JDK 21'),
    ('RAM', '2 GB', '4 GB atau lebih'),
    ('Penyimpanan', '10 GB', '50 GB atau lebih'),
    ('Port', '8080, 50000', '8080, 50000'),
]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
add_caption('Tabel 1. Spesifikasi sistem untuk instalasi Jenkins pada Windows')

doc.add_heading('3.3 Langkah-Langkah Instalasi Jenkins di Windows', level=2)
add_paragraph(
    'Instalasi Jenkins di Windows meliputi beberapa tahapan: instalasi JDK, pengunduhan installer MSI Jenkins, konfigurasi setup wizard, dan konfigurasi awal melalui browser.',
)
doc.add_picture('images/diagram_installation.png', width=Inches(6))
add_caption('Gambar 2. Tahapan instalasi Jenkins di Windows dari JDK hingga konfigurasi awal')

doc.add_heading('3.3.1 Instalasi JDK dan Konfigurasi JAVA_HOME', level=3)
add_number('Unduh JDK 17 atau 21 dari Eclipse Temurin atau distribusi OpenJDK lain yang didukung.')
add_number('Instal JDK ke folder C:\\Program Files\\Java\\jdk-17.x.x atau lokasi lain.')
add_number('Atur variabel lingkungan JAVA_HOME ke lokasi instalasi JDK dan tambahkan %JAVA_HOME%\\bin ke PATH.')

doc.add_heading('3.3.2 Instalasi Jenkins Windows MSI', level=3)
add_bullet('Unduh file installer Jenkins (.msi) dari https://www.jenkins.io/download.')
add_bullet('Jalankan installer dan pilih direktori instalasi Jenkins, port HTTP default 8080, serta konfigurasi layanan Windows.')
add_bullet('Pastikan Jenkins dapat mendeteksi instalasi Java secara otomatis selama setup wizard.')

doc.add_heading('3.3.3 Konfigurasi Awal Jenkins', level=3)
add_bullet('Akses Jenkins melalui browser dengan URL http://localhost:8080.')
add_bullet('Gunakan initialAdminPassword yang tersedia di C:\\ProgramData\\Jenkins\\.jenkins\\secrets\\initialAdminPassword.')
add_bullet('Pilih instalasi plugin yang disarankan untuk mempercepat konfigurasi awal.')


doc.add_heading('3.4 Konfigurasi Pipeline CI/CD di Windows', level=2)
add_paragraph(
    'Pipeline Jenkins di Windows menggunakan perintah bat untuk menjalankan langkah build dan test. Jenkinsfile yang ditulis sebagai Pipeline as Code memungkinkan definisi alur CI/CD disimpan bersama kode sumber.',
)
doc.add_picture('images/diagram_pipeline.png', width=Inches(6))
add_caption('Gambar 3. Alur pipeline CI/CD Jenkins di Windows meliputi checkout, build, test, dan deploy')

table = doc.add_table(rows=1, cols=4)
for idx, heading in enumerate(['No.', 'Tahap', 'Deskripsi', 'Contoh Perintah Windows']):
    cell = table.rows[0].cells[idx]
    cell.text = heading
for row_data in [
    ('1', 'Checkout', 'Mengambil kode terbaru dari repositori Git.', 'bat "git checkout main && git pull"'),
    ('2', 'Build', 'Membangun aplikasi menggunakan tools build seperti Maven atau Gradle.', 'bat "mvn package"'),
    ('3', 'Test', 'Menjalankan unit test dan pemeriksaan kualitas kode.', 'bat "mvn test"'),
    ('4', 'Deploy', 'Mendistribusikan artefak ke lingkungan staging atau produksi.', 'bat "deploy-staging.bat"'),
]:
    row_cells = table.add_row().cells
    row_cells[0].text = row_data[0]
    row_cells[1].text = row_data[1]
    row_cells[2].text = row_data[2]
    row_cells[3].text = row_data[3]
add_caption('Tabel 2. Tahapan pipeline CI/CD Jenkins di Windows beserta contoh perintah bat')

doc.add_heading('3.5 Plugin dan Integrasi Git', level=2)
add_paragraph(
    'Plugin penting yang direkomendasikan meliputi Git Plugin, Pipeline, Credentials, dan Blue Ocean. Plugin ini membantu Jenkins terintegrasi dengan repositori Git dan mendukung pipeline visual.',
)
add_bullet('Git Plugin untuk sinkronisasi dengan repositori Git dan GitHub.')
add_bullet('Pipeline untuk menjalankan Jenkinsfile dan mendefinisikan tahap CI/CD.')
add_bullet('Credentials untuk menyimpan token dan password secara aman.')
add_bullet('Blue Ocean untuk tampilan pipeline yang lebih ramah pengguna.')


doc.add_heading('BAB IV\nKESIMPULAN DAN SARAN', level=1)
doc.add_heading('4.1 Kesimpulan', level=2)
add_paragraph(
    'Jenkins adalah tools penting untuk mengotomasi proses CI/CD pada platform Windows. Dengan instalasi JDK, setup Windows MSI, dan konfigurasi pipeline, Jenkins dapat mempercepat proses build, test, dan deployment dengan cara yang repeatable dan terkontrol.',
)


doc.add_heading('4.2 Saran', level=2)
add_bullet('Selalu perbarui Jenkins dan plugin secara berkala untuk menjaga keamanan dan stabilitas.')
add_bullet('Gunakan backup konfigurasi Jenkins jika digunakan pada lingkungan produksi.')
add_bullet('Kembangkan Jenkinsfile dengan perintah bat yang jelas untuk meminimalkan kesalahan pada Windows.')


doc.add_heading('DAFTAR PUSTAKA', level=1)
references = [
    'Yulianto, A., & Nugraha, R. (2024). Implementasi Jenkins untuk Continuous Integration pada Sistem Operasi Windows. Jurnal Teknologi Informasi, 12(1), 45-54.',
    'Pratama, G., & Susilo, T. (2023). Peran Jenkins dalam Otomatisasi CI/CD pada Lingkungan Microsoft Windows. Journal of Software Engineering, 9(2), 120-128.',
    'Hartono, D., & Irawan, B. (2022). Analisis Penerapan DevOps dengan Jenkins di Platform Windows. Jurnal Komputer dan Informatika, 7(3), 210-218.',
    'Nugroho, E., & Putri, S. (2025). Studi Kasus Pipeline CI/CD Menggunakan Jenkins dan GitHub di Lingkungan Windows. International Journal of Computer Science, 15(4), 77-89.',
]
for ref in references:
    add_paragraph(ref)

# Save document
output_path = os.path.join(os.getcwd(), OUTPUT_DOCX)
doc.save(output_path)
print(f'DOCX created at {output_path}')
