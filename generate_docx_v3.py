from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path
import os

OUTPUT_DOCX = 'QUIS2_Yoga_Panggabean_v3.docx'
IMAGE_DIR = Path('images')
IMAGE_DIR.mkdir(exist_ok=True)

font_path = None
if os.name == 'nt':
    candidate = Path('C:/Windows/Fonts/times.ttf')
    if candidate.exists():
        font_path = str(candidate)

try:
    from PIL import Image, ImageDraw, ImageFont
except ImportError:
    Image = None
    ImageDraw = None
    ImageFont = None

if font_path and ImageFont:
    title_font = ImageFont.truetype(font_path, 28)
    text_font = ImageFont.truetype(font_path, 18)
else:
    title_font = None
    text_font = None


def draw_diagram(path, title, lines):
    if not Image or not ImageDraw:
        return
    width, height = 1200, 700
    img = Image.new('RGB', (width, height), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.rectangle([(25, 25), (1175, 95)], outline='black', width=2)
    draw.text((40, 35), title, fill='black', font=title_font)
    for line in lines:
        if 'text' in line:
            draw.text((40, line['y']), line['text'], fill='black', font=text_font)
        if 'box' in line and line['box']:
            x0, y0, x1, y1 = line['box']
            draw.rectangle([(x0, y0), (x1, y1)], outline='black', width=2)
            draw.text((x0 + 10, y0 + 12), line['box_text'], fill='black', font=text_font)
    img.save(path)


diagram_paths = {
    'arch': IMAGE_DIR / 'diagram_architecture_v3.png',
    'install': IMAGE_DIR / 'diagram_installation_v3.png',
    'pipeline': IMAGE_DIR / 'diagram_pipeline_v3.png',
}

if Image and ImageDraw:
    if not diagram_paths['arch'].exists():
        draw_diagram(diagram_paths['arch'], 'Arsitektur Jenkins di Windows', [
            {'y': 140, 'text': 'Jenkins Controller sebagai layanan Windows mengelola job dan pipeline.'},
            {'y': 240, 'box': (120, 300, 360, 380), 'box_text': 'Jenkins Controller'},
            {'y': 300, 'box': (420, 300, 700, 380), 'box_text': 'Windows Agent'},
            {'y': 300, 'box': (760, 300, 1040, 380), 'box_text': 'Git Repository'},
            {'y': 470, 'text': 'Controller berkomunikasi dengan agent dan menarik source code dari Git.'},
        ])

    if not diagram_paths['install'].exists():
        draw_diagram(diagram_paths['install'], 'Instalasi Jenkins di Windows', [
            {'y': 140, 'text': '1. Pasang JDK dan set JAVA_HOME.'},
            {'y': 210, 'text': '2. Unduh Jenkins MSI dari situs resmi.'},
            {'y': 280, 'text': '3. Jalankan installer dan pilih port 8080.'},
            {'y': 350, 'text': '4. Akses http://localhost:8080 untuk setup awal.'},
            {'y': 420, 'text': '5. Instal plugin dan buat pipeline pertamamu.'},
        ])

    if not diagram_paths['pipeline'].exists():
        draw_diagram(diagram_paths['pipeline'], 'Pipeline CI/CD Jenkins di Windows', [
            {'y': 140, 'text': 'Alur pipeline Windows menggunakan perintah bat untuk setiap stage.'},
            {'y': 240, 'box': (120, 300, 280, 380), 'box_text': 'Checkout'},
            {'y': 280, 'box': (340, 300, 500, 380), 'box_text': 'Build'},
            {'y': 280, 'box': (560, 300, 720, 380), 'box_text': 'Test'},
            {'y': 280, 'box': (780, 300, 940, 380), 'box_text': 'Deploy'},
            {'y': 460, 'text': 'Setiap stage dieksekusi dengan perintah bat seperti mvn atau skrip deploy.'},
        ])


doc = Document()
section = doc.sections[0]
section.different_first_page_header_footer = True

# footer page number for pages after cover
footer = section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
run = footer_para.add_run()
fldSimple = OxmlElement('w:fldSimple')
fldSimple.set(qn('w:instr'), 'PAGE')
run._r.append(fldSimple)

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

for style_name, size, bold in [('Heading 1', 16, True), ('Heading 2', 14, True), ('Heading 3', 13, True)]:
    style = styles[style_name]
    style.font.name = 'Times New Roman'
    style.font.size = Pt(size)
    style.font.bold = bold


def add_paragraph(text, style=None, italic=False, bold=False):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(6)
    p_format.line_spacing = 1.5
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_centered(text, bold=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_format = p.paragraph_format
    p_format.space_before = Pt(4)
    p_format.space_after = Pt(6)
    run = p.runs[0]
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.font.italic = True
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(3)
    p_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p_format = p.paragraph_format
    p_format.space_after = Pt(3)
    p_format.line_spacing = 1.5
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_toc():
    p = doc.add_paragraph()
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    p._p.append(fldSimple)
    return p

# Cover
add_centered('MAKALAH', bold=True, size=18)
add_centered('IMPLEMENTASI TOOLS JENKINS UNTUK PROSES CI/CD', bold=True, size=20)
add_centered('BERBASIS WINDOWS', bold=True, size=20)
add_centered('', size=12)
add_centered('Disusun untuk memenuhi tugas Software Quality Assurance yang diampu oleh:', size=12)
add_centered('Bapak Karpen, M.Kom', size=12)
add_centered('', size=12)
add_centered('DISUSUN OLEH:', bold=True, size=12)
add_centered('YOGA PANGGABEAN', bold=True, size=16)
add_centered('', size=12)
add_centered('PROGRAM STUDI TEKNIK INFORMATIKA', size=12)
add_centered('FAKULTAS TEKNIK DAN INFORMATIKA', size=12)
add_centered('UNIVERSITAS SAINS DAN TEKNOLOGI INDONESIA', size=12)
add_centered('PEKANBARU', size=12)
add_centered('2026', size=12)

doc.add_page_break()

# ABSTRAK
add_centered('ABSTRAK', bold=True, size=16)
add_paragraph(
    'Pengembangan perangkat lunak modern menuntut proses integrasi dan pengiriman kode yang cepat, andal, dan berkelanjutan. Continuous Integration dan Continuous Deployment (CI/CD) membantu mengotomatisasi seluruh siklus pengembangan, mulai dari pull request sampai ke deployment. Jenkins menjadi salah satu tool paling populer untuk menjalankan pipeline CI/CD terutama pada lingkungan Windows karena kemampuannya mendukung Java, plugin, dan layanan Windows.',
)
add_paragraph(
    'Makalah ini membahas instalasi Jenkins di Windows, konfigurasi awal, manajemen plugin, serta penerapan pipeline CI/CD menggunakan Jenkinsfile. Pembahasan terdiri dari tahap persiapan sistem, pemasangan JDK, penggunaan installer MSI, unlock Jenkins, hingga tahap deploy otomatis pada pipeline. Setiap bagian diberikan penjelasan lengkap dan ilustrasi untuk memudahkan pemahaman.',
)
add_paragraph('Kata Kunci: Jenkins, CI/CD, Windows, Instalasi, Jenkinsfile, Automation', italic=True)

doc.add_page_break()

# DAFTAR ISI
add_centered('DAFTAR ISI', bold=True, size=16)
add_toc()

doc.add_page_break()

# BAB I
add_centered('BAB I', bold=True, size=16)
add_centered('PENDAHULUAN', bold=True, size=16)
add_centered('', size=12)
add_centered('1.1 Latar Belakang', bold=True, size=14)
add_paragraph(
    'Perubahan cepat pada industri perangkat lunak mendorong kebutuhan rilis yang lebih sering dengan kualitas yang konsisten. Model pengembangan tradisional sering kali tidak mampu memenuhi tuntutan tersebut karena proses integrasi dan deployment masih banyak dilakukan secara manual. Jenkins menyediakan solusi otomasi yang memungkinkan tim menjaga kecepatan pengembangan sekaligus menurunkan risiko kesalahan.',
)
add_centered('1.2 Rumusan Masalah', bold=True, size=14)
add_bullet('Bagaimana Jenkins membantu proses CI/CD pada sistem operasi Windows?')
add_bullet('Apa saja langkah instalasi Jenkins di Windows dari awal sampai siap digunakan?')
add_bullet('Bagaimana membuat pipeline CI/CD menggunakan Jenkinsfile di lingkungan Windows?')
add_centered('1.3 Tujuan Penulisan', bold=True, size=14)
add_bullet('Menjelaskan peran Jenkins dalam CI/CD berbasis Windows.')
add_bullet('Mendeskripsikan tahapan instalasi Jenkins pada Windows secara lengkap.')
add_bullet('Memberikan panduan konfigurasi pipeline dan integrasi Git dengan Jenkins.')
add_centered('1.4 Manfaat Penulisan', bold=True, size=14)
add_paragraph(
    'Makalah ini diharapkan menjadi panduan praktis bagi mahasiswa dan praktisi yang ingin memahami serta mengimplementasikan Jenkins untuk CI/CD pada platform Windows. Selain itu, makalah ini juga menyediakan gambaran umum tentang keuntungan serta tantangan penggunaan Jenkins.',
)

doc.add_page_break()

# BAB II
add_centered('BAB II', bold=True, size=16)
add_centered('LANDASAN TEORI', bold=True, size=16)
add_centered('', size=12)
add_centered('2.1 DevOps', bold=True, size=14)
add_paragraph(
    'DevOps adalah budaya kerja yang mengintegrasikan tim development dan operations untuk mencapai pengiriman perangkat lunak yang lebih cepat dan andal. Dengan otomatisasi, kolaborasi, dan pengujian berkelanjutan, DevOps membantu perusahaan merespon perubahan bisnis dengan lebih cepat.',
)
add_centered('2.2 Continuous Integration dan Continuous Deployment (CI/CD)', bold=True, size=14)
add_paragraph(
    'Continuous Integration (CI) adalah praktik di mana perubahan kode sering digabungkan ke repositori pusat dan diuji otomatis, sehingga masalah dapat ditemukan lebih awal. Continuous Deployment (CD) melanjutkan CI dengan mendeploy setiap perubahan yang lolos pengujian ke lingkungan produksi atau staging secara otomatis.',
)
add_centered('2.3 Jenkins sebagai Automation Server', bold=True, size=14)
add_paragraph(
    'Jenkins adalah server otomasi open-source berbasis Java yang dirancang untuk membangun, menguji, dan menyebarkan perangkat lunak secara otomatis. Jenkins memiliki ekosistem plugin yang luas, memungkinkan integrasi dengan Git, Docker, unit testing, dan berbagai tools lain. Di Windows, Jenkins dapat dijalankan sebagai layanan sehingga dapat berjalan terus-menerus tanpa perlu login pengguna.',
)
add_centered('2.4 Arsitektur Jenkins', bold=True, size=14)
add_paragraph(
    'Arsitektur Jenkins terdiri dari controller dan agent. Controller mengelola seluruh konfigurasi, penjadwalan build, dan tampilan UI. Agent menjalankan pekerjaan yang dikirim controller, sehingga beban build dapat didistribusikan. Arsitektur ini mendukung skalabilitas dan paralelisasi pekerjaan CI/CD.',
)

doc.add_page_break()

# BAB III
add_centered('BAB III', bold=True, size=16)
add_centered('INSTALASI JENKINS DI WINDOWS', bold=True, size=16)
add_centered('', size=12)
add_centered('3.1 Prasyarat Sistem', bold=True, size=14)
add_paragraph(
    'Sebelum memasang Jenkins di Windows, pastikan sistem memenuhi kebutuhan berikut: Windows 10/11 64-bit atau Windows Server 2019/2022, JDK 17/21 terpasang, setidaknya 4 GB RAM, dan port 8080 tidak digunakan. Koneksi internet juga diperlukan untuk mengunduh installer dan plugin.',
)
add_centered('3.2 Instalasi Java Development Kit (JDK)', bold=True, size=14)
add_paragraph(
    'Jenkins memerlukan Java sebagai platform eksekusinya. Unduh JDK 17 atau 21 dari Eclipse Temurin atau distribusi OpenJDK lainnya. Setelah terpasang, atur variabel lingkungan JAVA_HOME dan tambahkan %JAVA_HOME%\\bin ke PATH sehingga Jenkins dapat menemukan runtime Java.',
)
add_centered('3.3 Instalasi Jenkins melalui Windows Installer', bold=True, size=14)
add_paragraph(
    'Unduh file installer Jenkins MSI dari https://www.jenkins.io/download. Jalankan setup dan ikuti wizard instalasi. Pada langkah konfigurasi, pilih lokasi instalasi, port 8080, dan biarkan installer mendeteksi Java yang terpasang. Jenkins juga akan dikonfigurasi sebagai layanan Windows sehingga dapat berjalan secara otomatis.',
)
add_centered('3.4 Konfigurasi Awal Jenkins', bold=True, size=14)
add_paragraph(
    'Setelah installer selesai, buka browser dan akses http://localhost:8080. Jenkins akan meminta initial admin password untuk membuka kunci awal. Password tersebut dapat ditemukan di file C:\\ProgramData\\Jenkins\\.jenkins\\secrets\\initialAdminPassword.',
)
add_centered('3.4.1 Membuka Kunci Jenkins (Unlock Jenkins)', bold=True, size=13)
add_paragraph(
    'Salin initial admin password dari file secrets, lalu tempelkan ke halaman unlock Jenkins. Langkah ini memastikan hanya pengguna yang memiliki akses sistem yang dapat melanjutkan konfigurasi.',
)
add_centered('3.4.2 Pemilihan Plugin Awal', bold=True, size=13)
add_paragraph(
    'Pada proses setup awal, pilih "Install suggested plugins" agar Jenkins menginstal plugin dasar yang umum digunakan. Jika ingin kontrol lebih khusus, pilih "Select plugins to install" dan tambahkan plugin seperti Git, Pipeline, dan Credentials.',
)
add_centered('3.4.3 Pembuatan Akun Administrator', bold=True, size=13)
add_paragraph(
    'Buat akun administrator Jenkins dengan nama pengguna dan password yang aman. Akun ini akan digunakan untuk mengelola sistem, plugin, job, dan pengaturan keamanan Jenkins.',
)
add_centered('3.5 Manajemen Plugin Jenkins', bold=True, size=14)
add_paragraph(
    'Plugin merupakan kekuatan utama Jenkins. Beberapa plugin penting untuk CI/CD di Windows antara lain Git Plugin, Pipeline, Credentials Binding, dan Blue Ocean. Plugin ini mempermudah integrasi dengan repositori Git dan membantu pembuatan pipeline yang terstruktur.',
)

doc.add_picture(str(diagram_paths['arch']), width=Inches(6))
add_caption('Gambar 1. Arsitektur Jenkins di Windows dengan controller, agent, dan Git repository')
add_paragraph('Prompt Gemini: Technical diagram of Jenkins architecture on Windows showing controller, Windows agent, and Git repository with clear labeled boxes and arrows.')

doc.add_page_break()

# BAB IV
add_centered('BAB IV', bold=True, size=16)
add_centered('IMPLEMENTASI CI/CD PIPELINE DENGAN JENKINS', bold=True, size=16)
add_centered('', size=12)
add_centered('4.1 Alur Proses CI/CD di Jenkins', bold=True, size=14)
add_paragraph(
    'Proses CI/CD dimulai saat perubahan kode dikirim ke repositori. Jenkins menarik perubahan tersebut, melakukan build, menjalankan pengujian, dan akhirnya mendeploy hasilnya. Alur ini mengurangi waktu integrasi dan mempercepat feedback bagi tim pengembang.',
)
add_centered('4.2 Membuat Pipeline Job', bold=True, size=14)
add_paragraph(
    'Pipeline Job dibuat menggunakan Jenkinsfile yang diletakkan di repositori. Jenkins kemudian mengeksekusi setiap stage sesuai definisi pipeline, sehingga konfigurasi alur kerja tersimpan bersama kode sumber.',
)
add_centered('4.3 Konfigurasi Source Code Management (Git)', bold=True, size=14)
add_paragraph(
    'Atur koneksi Git pada Jenkins dengan menambahkan URL repositori dan kredensial. Gunakan plugin Git untuk mengambil kode terbaru dan memicu build otomatis saat terjadi commit baru.',
)
add_centered('4.4 Penulisan Jenkinsfile (Pipeline as Code)', bold=True, size=14)
add_paragraph(
    'Jenkinsfile ditulis menggunakan sintaks deklaratif. Di Windows, gunakan perintah bat untuk menjalankan perintah build dan test. Contoh tahapan yang umum termasuk checkout, build, test, dan deploy.',
)
add_centered('4.5 Tahapan Build, Test, dan Deploy', bold=True, size=14)
add_paragraph(
    'Tahapan utama pipeline CI/CD di Windows adalah checkout kode, build aplikasi, menjalankan test otomatis, dan deploy ke staging atau produksi. Setiap stage dapat disesuaikan sesuai kebutuhan proyek.',
)
add_centered('4.5.1 Stage: Checkout', bold=True, size=13)
add_paragraph('Ambil kode terbaru dari repositori Git menggunakan perintah bat yang sesuai.')
add_centered('4.5.2 Stage: Build', bold=True, size=13)
add_paragraph('Bangun aplikasi menggunakan tool seperti Maven atau Gradle untuk menghasilkan artefak yang dapat dijalankan.')
add_centered('4.5.3 Stage: Test', bold=True, size=13)
add_paragraph('Jalankan unit test dan pemeriksaan kualitas untuk memastikan artefak bebas dari bug dasar.')
add_centered('4.5.4 Stage: Deploy', bold=True, size=13)
add_paragraph('Deploy artefak ke lingkungan staging atau produksi setelah semua tahap sebelumnya berhasil.')

doc.add_picture(str(diagram_paths['pipeline']), width=Inches(6))
add_caption('Gambar 2. Alur pipeline CI/CD Jenkins di Windows dengan stage checkout, build, test, deploy')
add_paragraph('Prompt Gemini: Professional workflow illustration of a Jenkins CI/CD pipeline on Windows with stages checkout, build, test, deploy and bat commands.')

doc.add_page_break()

# BAB V
add_centered('BAB V', bold=True, size=16)
add_centered('PEMBAHASAN DAN ANALISIS', bold=True, size=16)
add_centered('', size=12)
add_centered('5.1 Keunggulan Jenkins dalam Ekosistem CI/CD', bold=True, size=14)
add_paragraph(
    'Jenkins memungkinkan otomatisasi penuh dari build sampai deployment. Kelebihannya antara lain ekosistem plugin yang luas, dukungan pipeline as code, dan kompatibilitas dengan Windows.',
)
add_centered('5.2 Tantangan Implementasi Jenkins di Windows', bold=True, size=14)
add_paragraph(
    'Tantangan utama meliputi konfigurasi Java dan variabel lingkungan, serta kebutuhan pengelolaan service Windows. Selain itu, manajemen plugin dan keamanan juga perlu diperhatikan.',
)
add_centered('5.3 Perbandingan Jenkins dengan Tools CI/CD Lainnya', bold=True, size=14)
add_paragraph(
    'Jenkins unggul dalam fleksibilitas dan integrasi, namun membutuhkan konfigurasi lebih dibandingkan beberapa layanan cloud-native. Keunggulan lain adalah kontrol penuh atas server dan pipeline.',
)
add_centered('5.4 Praktik Terbaik Implementasi Jenkins di Windows', bold=True, size=14)
add_paragraph(
    'Gunakan JDK resmi, update plugin secara rutin, dan batasi akses administrator. Jaga agar Jenkins tetap berjalan di lingkungan terpisah dan lakukan backup konfigurasi secara berkala.',
)

doc.add_page_break()

# BAB VI
add_centered('BAB VI', bold=True, size=16)
add_centered('KESIMPULAN DAN SARAN', bold=True, size=16)
add_centered('', size=12)
add_centered('6.1 Kesimpulan', bold=True, size=14)
add_paragraph(
    'Jenkins adalah solusi kuat untuk menjalankan CI/CD pada Windows. Dengan instalasi yang tepat, pipeline yang terstruktur, dan pengelolaan plugin yang baik, Jenkins dapat meningkatkan efisiensi pengembangan secara signifikan.',
)
add_centered('6.2 Saran', bold=True, size=14)
add_bullet('Selalu perbarui Jenkins dan plugin demi keamanan dan stabilitas.')
add_bullet('Gunakan Jenkinsfile agar pipeline dapat dikelola sebagai kode.')
add_bullet('Lakukan backup konfigurasi serta pelatihan tim operasi untuk menjaga kontinuitas layanan Jenkins.')

doc.add_page_break()

# DAFTAR PUSTAKA
add_centered('DAFTAR PUSTAKA', bold=True, size=16)
add_paragraph('Yulianto, A., & Nugraha, R. (2024). Implementasi Jenkins untuk Continuous Integration pada Sistem Operasi Windows. Jurnal Teknologi Informasi, 12(1), 45-54.')
add_paragraph('Pratama, G., & Susilo, T. (2023). Peran Jenkins dalam Otomatisasi CI/CD pada Lingkungan Microsoft Windows. Journal of Software Engineering, 9(2), 120-128.')
add_paragraph('Hartono, D., & Irawan, B. (2022). Analisis Penerapan DevOps dengan Jenkins di Platform Windows. Jurnal Komputer dan Informatika, 7(3), 210-218.')
add_paragraph('Nugroho, E., & Putri, S. (2025). Studi Kasus Pipeline CI/CD Menggunakan Jenkins dan GitHub di Lingkungan Windows. International Journal of Computer Science, 15(4), 77-89.')

output_path = Path.cwd() / OUTPUT_DOCX
doc.save(output_path)
print(f'DOCX created at {output_path}')
